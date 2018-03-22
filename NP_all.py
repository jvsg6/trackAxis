#!/usr/bin/python2.7
# -*- coding: utf-8 -*-
import sys
import os

import struct
import openpyxl
import docx
import geopy
import argparse, arghelper

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from geographiclib.geodesic import Geodesic
from geopy.distance import vincenty
from grids.gridNostraResults import *
from additionalfunctions import make_sure_path_exists
from math import cos,sin,radians,sqrt,fabs,acos,pi,degrees,floor,ceil
from openpyxl import load_workbook
sourceLocation = geopy.Point(36.1449943051, 33.5403694792) #lat lon #АЭС Аккую

pointsForAnalysis = []
arrayOfResultsForPoints = []
nameConfInt = ""
tablesCounter = 0
shift = 0

def reaNPFromFile(pathToFile):
	inp = open(pathToFile,"r")
	lines = inp.readlines()
	inp.close()
	lines = [line for line in lines if line.find("#") < 0]
	pointsForAnalysis = []
	for line in lines:
		pointsForAnalysis.append([str(x.strip("\"")) if i == 0 else float(x) for (i,x) in enumerate(filter(None,line.split(",")))])
	
	idForSort = -1
	for pointId, point in enumerate(pointsForAnalysis):
		idForSort = len(point)
		pointsForAnalysis[pointId].append(vincenty(sourceLocation, geopy.Point(point[2], point[1]) , ellipsoid='WGS-84').km)
	if(idForSort>0):
		pointsForAnalysis = sorted(pointsForAnalysis, key=lambda point: point[idForSort])
		
	return pointsForAnalysis


def sumGridForDose(*args):
	if len(args) == 0:
		return None
	result = [[0.0 for j in range(len(arrayOfResultsForPoints[i][0]))] for i in range(len(pointsForAnalysis))]
	for arg in args:
		for pointId in range(len(pointsForAnalysis)):
			for j in range(len(arg[pointId])):
				result[pointId][j] = result[pointId][j] + arg[pointId][j]
	for arg in args:
		del arg[:]
	return result
def sumGridForDoseNew(*args):
	if len(args) == 0:
		return None
	result = [[0.0 for j in range(len(arrayOfResultsForPoints[i][0]))] for i in range(len(pointsForAnalysis))]
	for arg in args:
		idx = int(arg[1:])
		for pointId in range(len(pointsForAnalysis)):
			for j in range(len(arrayOfResultsForPoints[pointId][idx])):
				result[pointId][j] = result[pointId][j] + arrayOfResultsForPoints[pointId][idx][j]
	return result
def maxGridForDoseNew(*args):
	if len(args) == 0:
		return None
	result = [[0.0 for j in range(len(arrayOfResultsForPoints[i][0]))] for i in range(len(pointsForAnalysis))]
	for arg in args:
		idx = int(arg[1:])
		for pointId in range(len(pointsForAnalysis)):
			for j in range(len(arrayOfResultsForPoints[pointId][idx])):
				if arrayOfResultsForPoints[pointId][idx][j] > result[pointId][j]:
					result[pointId][j] = arrayOfResultsForPoints[pointId][idx][j]
	return result
def prepareArrayForVarSeries():
	global pointsForAnalysis, arrayOfResultsForPoints
	d = iFileFolder
	listOfFiles = [os.path.join(d,o) for o in os.listdir(d) if (os.path.isfile(os.path.join(d,o)) and os.path.join(d,o).find(".bin")>=0)]
	for k in range(len(listOfFiles)): 
		fileName = iFileFolder+str("/f"+str(k)+".bin")
		inp = open(fileName, 'rb')
		datastr = inp.read(64)
		offsetStart = 64
		headerStr = list(struct.unpack('<%dd' % (len(datastr)/8), datastr))
		grid = GNR_Base()
		grid.initByParameters(int(headerStr[0]),int(headerStr[1]),headerStr[2],headerStr[3],headerStr[4],headerStr[5],headerStr[6],headerStr[7])
		del grid.data
		countx = int(headerStr[0])
		county = int(headerStr[1])
		offsetForNextGrid = countx*county*4
		gridsCount = (os.stat(fileName).st_size-offsetStart)/offsetForNextGrid-1
		for gridId in range(gridsCount):
			offset = offsetStart + gridId*offsetForNextGrid
			inp.seek(offset)
			grid.data = struct.unpack('<%df' % (countx*county), inp.read(offsetForNextGrid))
			for pointId,point in enumerate(pointsForAnalysis):
				if gridId == 0:
					arrayOfResultsForPoints[pointId].append([])
				arrayOfResultsForPoints[pointId][k].append(grid.getValue(point[1],point[2]))
			del grid.data
		inp.close()
		#if k == 30:
		#	break
		print "done", str(float(k)/len(listOfFiles)*100.0)+" %"
		
	#вариант Асфандиярова (не совсем корректный)
	#for pointId,point in enumerate(pointsForAnalysis):
	#	for k in range(len(arrayOfResultsForPoints[pointId])):
	#		arrayOfResultsForPoints[pointId][k].sort()
	return
	
def getValueForConfidence(pointId,array,confVal):
	idx = int(ceil(float(len(array[pointId]))*confVal/100.0))-1
	tmpArr = array[pointId][:]
	tmpArr.sort()
	val = tmpArr[idx]
	del tmpArr
	return val

def prepareNew(idf):
	if len(idf) == 0:
		return None
	result = [[0.0 for j in range(len(arrayOfResultsForPoints[i][0]))] for i in range(len(pointsForAnalysis))]
	idx = int(idf[1:])
	for pointId in range(len(pointsForAnalysis)):
		for j in range(len(arrayOfResultsForPoints[pointId][idx])):
			result[pointId][j] = result[pointId][j] + arrayOfResultsForPoints[pointId][idx][j]
	return result

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def addTextToTablesHeaderCell(cell,text,alignment = WD_ALIGN_PARAGRAPH.LEFT):
	p1 = cell.add_paragraph(style='name-of-t-style')
   	p1.alignment = alignment
	run = p1.add_run(text, style = 'name-of-ch_t-style')
	run.bold = True

def addTextToTablesCell(cell,text,alignment = WD_ALIGN_PARAGRAPH.LEFT):
	p1 = cell.add_paragraph(style='name-of-t-style')
   	p1.alignment = alignment
	p1.add_run(text, style = 'name-of-ch_t-style')

def addSingleTable(document,scenario,mainTitle,multiplier,separate, *args):
	if separate:
			document = setDocumentStyle()
	global tablesCounter	
	tablesCounter = tablesCounter + 1
	p = document.add_paragraph('', style='name-of-p-style')
	p.add_run(u'Таблица {} – '.format(str(tablesCounter))+mainTitle+". "+scenario, style = 'name-of-ch-style')

	table = document.add_table(rows=1, cols=len(args)+1, style='Table Grid')
	hdr_cells = table.rows[0].cells
	delete_paragraph(hdr_cells[0].paragraphs[-1])
	addTextToTablesHeaderCell(hdr_cells[0], u'Населенные пункты (долгота; широта; расстояние, км)', WD_ALIGN_PARAGRAPH.CENTER)
	for pos,arg in enumerate(args):
		delete_paragraph(hdr_cells[pos+1].paragraphs[-1])
		addTextToTablesHeaderCell(hdr_cells[pos+1], arg[1], WD_ALIGN_PARAGRAPH.CENTER)

	countF = len(args)	
	vals = []
	for pos,arg in enumerate(args):
		vals.append(prepareNew(arg[0]))
	row_cells = []
	for pointId, point in enumerate(pointsForAnalysis):
		for iid, i in enumerate(listOfConInt):
			#print pointId, iid 
			row_cells.append(table.add_row().cells)
		#print row_cells
		col_cells = table.column_cells(1)
		a = table.cell(pointId*len(listOfConInt)+1, 0)
		b = table.cell((pointId+1)*len(listOfConInt), 0)
		#print pointId*len(listOfConInt), (pointId+1)*len(listOfConInt)
		A = a.merge(b)
		delete_paragraph(row_cells[pointId*len(listOfConInt)][0].paragraphs[-1])		
		addTextToTablesCell(row_cells[pointId*len(listOfConInt)][0],point[0]+" ("+ format(point[1], '.6f').replace('.', ',')+"; "+format(point[2], '.6f').replace('.', ',')+"; "+format(point[3], '.1f').replace('.', ',')+")")
		#print "fffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffff"
		for pos,val in enumerate(vals):
			
			for jjd, j in enumerate(listOfConInt):	
				delete_paragraph(row_cells[pointId*len(listOfConInt)+jjd][pos+1].paragraphs[-1])
				addTextToTablesCell(row_cells[pointId*len(listOfConInt)+jjd][pos+1],format(getValueForConfidence(pointId,val, float(listOfConInt[jjd]))*multiplier, '.2E').replace('.', ','), WD_ALIGN_PARAGRAPH.CENTER)
	
	document.add_page_break()	
	if separate == True:
		document.save(oFileFolder+'/appendix1_{}.docx'.format(tablesCounter))
	return 

def addComplexTable(document,scenario,mainTitle,multiplier,separate, t1,t2,t3,t4,t5,t6,t7,t8):
	if separate:
			document = setDocumentStyle()
	global tablesCounter	
	tablesCounter = tablesCounter + 1
	p = document.add_paragraph('', style='name-of-p-style')
	p.add_run(u'Таблица {} – '.format(str(tablesCounter))+mainTitle+". "+scenario, style = 'name-of-ch-style')

	table = document.add_table(rows=2, cols=9, style='Table Grid')
	a1 = table.cell(0, 0)
	b1 = table.cell(1, 0)
	A1 = a1.merge(b1)
	delete_paragraph(A1.paragraphs[-1])
	addTextToTablesHeaderCell(A1, u'Населенные пункты (долгота; широта; расстояние, км)', WD_ALIGN_PARAGRAPH.CENTER)

	a1 = table.cell(0, 1)
	b1 = table.cell(0, 2)
	A2 = a1.merge(b1)
	delete_paragraph(A2.paragraphs[-1])
	addTextToTablesHeaderCell(A2, u'Внешнее', WD_ALIGN_PARAGRAPH.CENTER)

	a1 = table.cell(0, 3)
	b1 = table.cell(0, 8)
	A3 = a1.merge(b1)
	delete_paragraph(A3.paragraphs[-1])
	addTextToTablesHeaderCell(A3, u'Ингаляция', WD_ALIGN_PARAGRAPH.CENTER)
	
	a = table.cell(1, 1)	
	delete_paragraph(a.paragraphs[-1])
	addTextToTablesHeaderCell(a, t1[1], WD_ALIGN_PARAGRAPH.CENTER)

	a = table.cell(1, 2)	
	delete_paragraph(a.paragraphs[-1])
	addTextToTablesHeaderCell(a, t2[1], WD_ALIGN_PARAGRAPH.CENTER)

	a = table.cell(1, 3)	
	delete_paragraph(a.paragraphs[-1])
	addTextToTablesHeaderCell(a, t3[1], WD_ALIGN_PARAGRAPH.CENTER)

	a = table.cell(1, 4)	
	delete_paragraph(a.paragraphs[-1])
	addTextToTablesHeaderCell(a, t4[1], WD_ALIGN_PARAGRAPH.CENTER)

	a = table.cell(1, 5)	
	delete_paragraph(a.paragraphs[-1])
	addTextToTablesHeaderCell(a, t5[1], WD_ALIGN_PARAGRAPH.CENTER)

	a = table.cell(1, 6)	
	delete_paragraph(a.paragraphs[-1])
	addTextToTablesHeaderCell(a, t6[1], WD_ALIGN_PARAGRAPH.CENTER)

	a = table.cell(1, 7)	
	delete_paragraph(a.paragraphs[-1])
	addTextToTablesHeaderCell(a, t7[1], WD_ALIGN_PARAGRAPH.CENTER)

	a = table.cell(1, 8)	
	delete_paragraph(a.paragraphs[-1])
	addTextToTablesHeaderCell(a, t8[1], WD_ALIGN_PARAGRAPH.CENTER)
	
	vals = []
	vals.append(prepareNew(t1[0]))
	vals.append(prepareNew(t2[0]))
	vals.append(prepareNew(t3[0]))
	vals.append(prepareNew(t4[0]))
	vals.append(prepareNew(t5[0]))
	vals.append(prepareNew(t6[0]))
	vals.append(prepareNew(t7[0]))
	vals.append(prepareNew(t8[0]))
	row_cells = []
	for pointId, point in enumerate(pointsForAnalysis):
		for iid, i in enumerate(listOfConInt):
			#print pointId, iid 
			row_cells.append(table.add_row().cells)
		#print row_cells
		col_cells = table.column_cells(0)
		a = table.cell(pointId*len(listOfConInt)+2, 0)
		b = table.cell((pointId+1)*len(listOfConInt)+1, 0)
		#print pointId*len(listOfConInt), (pointId+1)*len(listOfConInt)
		A = a.merge(b)
		delete_paragraph(row_cells[pointId*len(listOfConInt)][0].paragraphs[-1])		
		addTextToTablesCell(row_cells[pointId*len(listOfConInt)][0],point[0]+" ("+ format(point[1], '.6f').replace('.', ',')+"; "+format(point[2], '.6f').replace('.', ',')+"; "+format(point[3], '.1f').replace('.', ',')+")")
		for pos,val in enumerate(vals):
			for jjd, j in enumerate(listOfConInt):
				delete_paragraph(row_cells[pointId*len(listOfConInt)+jjd][pos+1].paragraphs[-1])
				addTextToTablesCell(row_cells[pointId*len(listOfConInt)+jjd][pos+1],format(getValueForConfidence(pointId,val, float(listOfConInt[jjd]))*multiplier, '.2E').replace('.', ','), WD_ALIGN_PARAGRAPH.CENTER)
	if separate == True:
		document.save(oFileFolder+'/appendix1_{}.docx'.format(tablesCounter))
	return 
def writeOrgans(ws, *args):
	global shift
	for nameId, name in enumerate(args):
		ws.cell(row = 3+shift, column = 2+nameId, value = name)
	return
# Таблица 1 – Прогнозируемые дозы за счет внешнего облучения от облака и от выпадений на поверхность земли (взвешенные по ОБЭ дозы за первые 10 ч), Гр . scenario
def createTable1(scenario, separate, createNew):
	global nameConfInt, shift
	_RED_MARROW = sumGridForDoseNew("f10","f11")
	_LUNGS = sumGridForDoseNew("f12","f13")
	_OVARIES = sumGridForDoseNew("f14","f15")
	_TESTES = sumGridForDoseNew("f16","f17")
	_THYROID = sumGridForDoseNew("f18","f19")
	_FOETUS = sumGridForDoseNew("f3","f7")
	
	wb = None
	if separate:
		wb = openpyxl.Workbook()
		ws = wb.worksheets[0]
		ws.title = u'Таблица 1'
		shift = 0
	else:
		if createNew:
			wb = openpyxl.Workbook()
			ws = wb.worksheets[0]
			ws.title = u'Таблицы'
		else:
			wb = load_workbook(filename = oFileFolder+"/Table.xlsx")
			ws = wb.worksheets[0]
			
	startCol = 1; endCol = 6; sratrRow = 1+shift; endRow = 1+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A1:F1')
	ws.cell(row = sratrRow, column = startCol,  value = u'Таблица 1 – Прогнозируемые дозы с уровнями доверия varSeries процентов за счет внешнего облучения от облака и от выпадений на поверхность земли (взвешенные по ОБЭ дозы за первые 10 ч), Гр. scenario'.replace("scenario",scenario).replace("varSeries", nameConfInt))
	
	startCol = 1; endCol = 1; sratrRow = 2+shift; endRow = 3+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A2:A3')
	ws.cell(row = sratrRow, column = startCol,  value = u'Населенные пункты (долгота; широта; расстояние, км)')
	
	startCol = 2; endCol = 6; sratrRow = 2+shift; endRow = 2+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('B2:F2')
	ws.cell(row = sratrRow, column = startCol,  value = u'Дозы на органы')
	
	
	ws.cell(row = 3+shift, column = 2, value = u'красный костный мозг')
	ws.cell(row = 3+shift, column = 3, value = u'легкие')
	ws.cell(row = 3+shift, column = 4, value = u'гонады')
	ws.cell(row = 3+shift, column = 5, value = u'щитовидная железа')
	ws.cell(row = 3+shift, column = 6, value = u'плод')
	
	organs = 5
	curPos = 4
	
	for pointId, point in enumerate(pointsForAnalysis):
		ws.cell(row = curPos+pointId*len(listOfConInt)+shift, column = 1, value = point[0]+" ("+ format(point[1], '.6f').replace('.', ',')+";"+format(point[2], '.6f').replace('.', ',')+";"+format(point[3], '.1f').replace('.', ',')+")")
		for  ConIntId, ConInt in enumerate(listOfConInt):
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=2, value=getValueForConfidence(pointId,_RED_MARROW, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=3, value=getValueForConfidence(pointId,_LUNGS, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=4, value=max(getValueForConfidence(pointId,_OVARIES, float(ConInt)),getValueForConfidence(pointId,_TESTES, float(listOfConInt[1]))))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=5, value=getValueForConfidence(pointId,_THYROID, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=6, value=getValueForConfidence(pointId,_FOETUS, float(ConInt)))	
	shift = shift + curPos+ConIntId+pointId*len(listOfConInt)+2

	for pointId, point in enumerate(pointsForAnalysis):
		for i in range(0, organs):
			ws.cell(row=curPos+pointId, column=2+i).number_format = '0.00E+00'
			
	if separate:		
		wb.save(oFileFolder+"/Table1.xlsx")
	else:
		wb.save(oFileFolder+"/Table.xlsx")
	
	del _RED_MARROW[:]
	del _LUNGS[:]
	del _OVARIES[:]
	del _TESTES[:]
	del _THYROID[:]
	del _FOETUS[:]
	
	return 

# Таблица 2 – Прогнозируемые дозы за счет: внутреннего облучения от ингаляционного поступления радионуклидов для взрослых (взвешенные по ОБЭ дозы по внутреннему пути облучения за период 30 суток), Гр . scenario
def createTable2(scenario, separate, createNew):
	global nameConfInt, shift
	_RED_MARROW = maxGridForDoseNew("f37")
	_THYROID = maxGridForDoseNew("f53")
	_LUNGS = maxGridForDoseNew("f45")
	_UPPER_LARGE_INTESTINE_WALL = maxGridForDoseNew("f61")
	_LOWER_LARGE_INTESTINE_WALL = maxGridForDoseNew("f69")
	_FOETUS = maxGridForDoseNew("f2")
	
	wb = None
	if separate:
		wb = openpyxl.Workbook()
		ws = wb.worksheets[0]
		ws.title = u'Таблица 2'
		shift = 0
	else:
		if createNew:
			wb = openpyxl.Workbook()
			ws = wb.worksheets[0]
		else:
			wb = load_workbook(filename = oFileFolder+"/Table.xlsx")
			ws = wb.worksheets[0]

	startCol = 1; endCol = 6; sratrRow = 1+shift; endRow = 1+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A1:F1')
	ws.cell(row = sratrRow, column = startCol,  value = u'Таблица 2 – Прогнозируемые дозы с уровнями доверия varSeries процентов за счет: внутреннего облучения от ингаляционного поступления радионуклидов для взрослых (взвешенные по ОБЭ дозы по внутреннему пути облучения за период 30 суток), Гр . scenario'.replace("scenario",scenario).replace("varSeries", nameConfInt))
	
	startCol = 1; endCol = 1; sratrRow = 2+shift; endRow = 3+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A2:A3')
	ws.cell(row = sratrRow, column = startCol,  value = u'Населенные пункты (долгота; широта; расстояние, км)')
	
	startCol = 2; endCol = 6; sratrRow = 2+shift; endRow = 2+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('B2:F2')
	ws.cell(row = sratrRow, column = startCol,  value = u'Дозы на органы (взрослые)')

	
	ws.cell(row = 3+shift, column = 2, value = u'красный костный мозг')
	ws.cell(row = 3+shift, column = 3, value = u'щитовидная железа')
	ws.cell(row = 3+shift, column = 4, value = u'легкие')
	ws.cell(row = 3+shift, column = 5, value = u'толстый кишечник')
	ws.cell(row = 3+shift, column = 6, value = u'плод')
	
	organs = 5
	curPos = 4
	
	for pointId, point in enumerate(pointsForAnalysis):
		ws.cell(row = curPos+pointId*len(listOfConInt)+shift, column = 1, value = point[0]+" ("+ format(point[1], '.6f').replace('.', ',')+";"+format(point[2], '.6f').replace('.', ',')+";"+format(point[3], '.1f').replace('.', ',')+")")
		for  ConIntId, ConInt in enumerate(listOfConInt):
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=2, value=getValueForConfidence(pointId,_RED_MARROW, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=3, value=getValueForConfidence(pointId,_THYROID, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=4, value=getValueForConfidence(pointId,_LUNGS, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=5, value=0.57*getValueForConfidence(pointId,_UPPER_LARGE_INTESTINE_WALL, float(ConInt)) + 0.43*getValueForConfidence(pointId,_LOWER_LARGE_INTESTINE_WALL, float(ConInt)) )
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=6, value=getValueForConfidence(pointId,_FOETUS, float(ConInt)))
	shift = shift + curPos+ConIntId+pointId*len(listOfConInt)+2
	
	for pointId, point in enumerate(pointsForAnalysis):
		for i in range(0, organs):
			ws.cell(row=curPos+pointId, column=2+i).number_format = '0.00E+00'
	
	if separate:		
		wb.save(oFileFolder+"/Table2.xlsx")
	else:
		wb.save(oFileFolder+"/Table.xlsx")
	
	del _RED_MARROW[:]
	del _THYROID[:]
	del _LUNGS[:]
	del _UPPER_LARGE_INTESTINE_WALL[:]
	del _LOWER_LARGE_INTESTINE_WALL[:]
	del _FOETUS[:]
	
	return

# Таблица 3 – Прогнозируемые дозы за счет: внутреннего облучения от ингаляционного поступления радионуклидов для детей (взвешенные по ОБЭ дозы по внутреннему пути облучения за период 30 суток), Гр . scenario
def createTable3(scenario, separate, createNew):
	global nameConfInt, shift
	_RED_MARROW = maxGridForDoseNew("f32","f33","f34","f35","f36")
	_THYROID = maxGridForDoseNew("f48","f49","f50","f51","f52")
	_LUNGS = maxGridForDoseNew("f40","f41","f42","f43","f44")
	_UPPER_LARGE_INTESTINE_WALL = maxGridForDoseNew("f56","f57","f58","f59","f60")
	_LOWER_LARGE_INTESTINE_WALL = maxGridForDoseNew("f64","f65","f66","f67","f68")
	
	wb = None
	if separate:
		wb = openpyxl.Workbook()
		ws = wb.worksheets[0]
		ws.title = u'Таблица 3'
		shift = 0
	else:
		if createNew:
			wb = openpyxl.Workbook()
			ws = wb.worksheets[0]
		else:
			wb = load_workbook(filename = oFileFolder+"/Table.xlsx")
			ws = wb.worksheets[0]
			
	startCol = 1; endCol = 5; sratrRow = 1+shift; endRow = 1+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A1:E1')
	ws.cell(row = sratrRow, column = startCol,  value = u'Таблица 3 – Прогнозируемые дозы с уровнями доверия varSeries процентов за счет: внутреннего облучения от ингаляционного поступления радионуклидов для детей (взвешенные по ОБЭ дозы по внутреннему пути облучения за период 30 суток), Гр . scenario'.replace("scenario",scenario).replace("varSeries", nameConfInt))
	
	startCol = 1; endCol = 1; sratrRow = 2+shift; endRow = 3+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A2:A3')
	ws.cell(row = sratrRow, column = startCol,  value = u'Населенные пункты (долгота; широта; расстояние, км)')
	
	startCol = 2; endCol = 5; sratrRow = 2+shift; endRow = 2+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('B2:E2')
	ws.cell(row = sratrRow, column = startCol,  value = u'Дозы на органы (дети)')

	writeOrgans(ws, u'красный костный мозг', u'щитовидная железа', u'легкие', u'толстый кишечник')

	organs = 4
	curPos = 4
	for pointId, point in enumerate(pointsForAnalysis):
		ws.cell(row = curPos+pointId*len(listOfConInt)+shift, column = 1, value = point[0]+" ("+ format(point[1], '.6f').replace('.', ',')+";"+format(point[2], '.6f').replace('.', ',')+";"+format(point[3], '.1f').replace('.', ',')+")")
		for  ConIntId, ConInt in enumerate(listOfConInt):
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=2, value=getValueForConfidence(pointId,_RED_MARROW, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=3, value=getValueForConfidence(pointId,_THYROID, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=4, value=getValueForConfidence(pointId,_LUNGS, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=5, value=0.57*getValueForConfidence(pointId,_UPPER_LARGE_INTESTINE_WALL, float(ConInt)) + 0.43*getValueForConfidence(pointId,_LOWER_LARGE_INTESTINE_WALL, float(ConInt)) )
	
	shift = shift + curPos+ConIntId+pointId*len(listOfConInt)+2
	for pointId, point in enumerate(pointsForAnalysis):
		for i in range(0, organs):
			ws.cell(row=curPos+pointId, column=2+i).number_format = '0.00E+00'
	
	if separate:		
		wb.save(oFileFolder+"/Table3.xlsx")
	else:
		wb.save(oFileFolder+"/Table.xlsx")
	
	del _RED_MARROW[:]
	del _THYROID[:]
	del _LUNGS[:]
	del _UPPER_LARGE_INTESTINE_WALL[:]
	del _LOWER_LARGE_INTESTINE_WALL[:]
	
	return

# Таблица 4 – Прогнозируемые дозы за счет: внешнего облучения от облака и от выпадений на поверхность земли (взвешенные по ОБЭ дозы за первые 30 суток), Гр . scenario
def createTable4(scenario, separate, createNew):
	global nameConfInt, shift
	_RED_MARROW = sumGridForDoseNew("f30","f31")
	_THYROID = sumGridForDoseNew("f46","f47")
	_LUNGS = sumGridForDoseNew("f38","f39")
	_UPPER_LARGE_INTESTINE_WALL = sumGridForDoseNew("f54","f55")
	_LOWER_LARGE_INTESTINE_WALL = sumGridForDoseNew("f62","f63")
	_FOETUS = sumGridForDoseNew("f3","f9")
	
	wb = None
	if separate:
		wb = openpyxl.Workbook()
		ws = wb.worksheets[0]
		ws.title = u'Таблица 4'
		shift = 0
	else:
		if createNew:
			wb = openpyxl.Workbook()
			ws = wb.worksheets[0]
		else:
			wb = load_workbook(filename = oFileFolder+"/Table.xlsx")
			ws = wb.worksheets[0]
	
	startCol = 1; endCol = 6; sratrRow = 1+shift; endRow = 1+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A1:F1')
	ws.cell(row = sratrRow, column = startCol,  value = u'Таблица 4 – Прогнозируемые дозы с уровнями доверия varSeries процентов за счет: внешнего облучения от облака и от выпадений на поверхность земли (взвешенные по ОБЭ дозы за первые 30 суток), Гр . scenario'.replace("scenario",scenario).replace("varSeries", nameConfInt))
	
	startCol = 1; endCol = 1; sratrRow = 2+shift; endRow = 3+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A2:A3')
	ws.cell(row = sratrRow, column = startCol,  value = u'Населенные пункты (долгота; широта; расстояние, км)')
	
	startCol = 2; endCol = 6; sratrRow = 2+shift; endRow = 2+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('B2:F2')
	ws.cell(row = sratrRow, column = startCol,  value = u'Дозы на органы')
	
	writeOrgans(ws, u'красный костный мозг', u'щитовидная железа', u'легкие', u'толстый кишечник', u'плод')
	
	organs = 5
	curPos = 4
	
	for pointId, point in enumerate(pointsForAnalysis):
		ws.cell(row = curPos+pointId*len(listOfConInt)+shift, column = 1, value = point[0]+" ("+ format(point[1], '.6f').replace('.', ',')+";"+format(point[2], '.6f').replace('.', ',')+";"+format(point[3], '.1f').replace('.', ',')+")")
		for  ConIntId, ConInt in enumerate(listOfConInt):
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=2, value=getValueForConfidence(pointId,_RED_MARROW, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=3, value=getValueForConfidence(pointId,_THYROID, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=4, value=getValueForConfidence(pointId,_LUNGS, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=5, value=0.57*getValueForConfidence(pointId,_UPPER_LARGE_INTESTINE_WALL, float(ConInt)) + 0.43*getValueForConfidence(pointId,_LOWER_LARGE_INTESTINE_WALL, float(ConInt)) )
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=6, value=getValueForConfidence(pointId,_FOETUS, float(ConInt)))
	shift = shift + curPos+ConIntId+pointId*len(listOfConInt)+2
	
	for pointId, point in enumerate(pointsForAnalysis):
		for i in range(0, organs):
			ws.cell(row=curPos+pointId, column=2+i).number_format = '0.00E+00'
	
	if separate:		
		wb.save(oFileFolder+"/Table4.xlsx")
	else:
		wb.save(oFileFolder+"/Table.xlsx")
		
	del _RED_MARROW[:]
	del _THYROID[:]
	del _LUNGS[:]
	del _UPPER_LARGE_INTESTINE_WALL[:]
	del _LOWER_LARGE_INTESTINE_WALL[:]
	del _FOETUS[:]
	
	return

# Таблица 5 – Прогнозируемые эквивалентные и эффективная дозы за счет: облучения от облака и от выпадений на поверхность земли и внутреннего облучения от ингаляционного поступления радионуклидов для взрослых (доза за 10 суток), Зв . scenario
def createTable5(scenario, separate, createNew):
	global nameConfInt, shift
	_RED_MARROW = sumGridForDoseNew("f142","f143","f149")
	_THYROID = sumGridForDoseNew("f166","f167","f173")
	_LUNGS = sumGridForDoseNew("f150","f151","f157")
	_SMALL_INTESTINE_WALL = sumGridForDoseNew("f190","f191","f197")
	_SKIN = sumGridForDoseNew("f158","f159","f165")
	_TESTES = sumGridForDoseNew("f182","f183","f189")
	_OVARIES = sumGridForDoseNew("f174","f175","f181")
	_EFFECTIVE_DOSE = sumGridForDoseNew("f134","f135","f141")
	
	wb = None
	if separate:
		wb = openpyxl.Workbook()
		ws = wb.worksheets[0]
		ws.title = u'Таблица 5'
		shift = 0
	else:
		if createNew:
			wb = openpyxl.Workbook()
			ws = wb.worksheets[0]
		else:
			wb = load_workbook(filename = oFileFolder+"/Table.xlsx")
			ws = wb.worksheets[0]

	
	startCol = 1; endCol = 8; sratrRow = 1+shift; endRow = 1+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A1:H1')
	ws.cell(row = sratrRow, column = startCol,  value = u'Таблица 5 – Прогнозируемые эквивалентные и эффективная дозы с уровнями доверия varSeries процентов за счет: облучения от облака и от выпадений на поверхность земли и внутреннего облучения от ингаляционного поступления радионуклидов для взрослых (доза за 10 суток), Зв . scenario'.replace("scenario",scenario).replace("varSeries", nameConfInt))
	
	startCol = 1; endCol = 1; sratrRow = 2+shift; endRow = 3+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A2:A3')
	ws.cell(row = sratrRow, column = startCol,  value = u'Населенные пункты (долгота; широта; расстояние, км)')
	
	startCol = 2; endCol = 8; sratrRow = 2+shift; endRow = 2+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('B2:H2')
	ws.cell(row = sratrRow, column = startCol,  value = u'Дозы на органы (взрослые)')
	
	writeOrgans(ws, u'красный костный мозг', u'щитовидная железа', u'легкие', u'тонкий кишечник',  u'кожа', u'гонады', u'эффективная доза')
	
	organs = 7
	curPos = 4
	
	for pointId, point in enumerate(pointsForAnalysis):
		ws.cell(row = curPos+pointId*len(listOfConInt)+shift, column = 1, value = point[0]+" ("+ format(point[1], '.6f').replace('.', ',')+";"+format(point[2], '.6f').replace('.', ',')+";"+format(point[3], '.1f').replace('.', ',')+")")
		for  ConIntId, ConInt in enumerate(listOfConInt):
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=2, value=getValueForConfidence(pointId,_RED_MARROW, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=3, value=getValueForConfidence(pointId,_THYROID, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=4, value=getValueForConfidence(pointId,_LUNGS, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=5, value=getValueForConfidence(pointId,_SMALL_INTESTINE_WALL, float(ConInt)) )
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=6, value=getValueForConfidence(pointId,_SKIN, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=7, value=max(getValueForConfidence(pointId,_OVARIES, float(ConInt)),getValueForConfidence(pointId,_TESTES, float(ConInt))))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=8, value=getValueForConfidence(pointId,_EFFECTIVE_DOSE, float(ConInt)) )
			
	shift = shift + curPos+ConIntId+pointId*len(listOfConInt)+2
	for pointId, point in enumerate(pointsForAnalysis):
		for i in range(0, organs):
			ws.cell(row=curPos+pointId, column=2+i).number_format = '0.00E+00'
	
	if separate:		
		wb.save(oFileFolder+"/Table5.xlsx")
	else:
		wb.save(oFileFolder+"/Table.xlsx")
	
	del _RED_MARROW[:]
	del _THYROID[:]
	del _LUNGS[:]
	del _SMALL_INTESTINE_WALL[:]
	del _SKIN[:]
	del _TESTES[:]
	del _OVARIES[:]
	del _EFFECTIVE_DOSE[:]
	
	return

# Таблица 6 – Прогнозируемые эквивалентные и эффективная дозы за счет: облучения от облака и от выпадений на поверхность земли и внутреннего облучения от ингаляционного поступления радионуклидов для детей (доза за 10 суток), Зв . scenario
def createTable6(scenario, separate, createNew):
	global nameConfInt, shift
	_RED_MARROW = sumGridForDose(sumGridForDoseNew("f142","f143"),maxGridForDoseNew("f144","f145","f146","f147","f148"))
	_THYROID = sumGridForDose(sumGridForDoseNew("f166","f167"),maxGridForDoseNew("f168","f169","f170","f171","f172"))
	_LUNGS = sumGridForDose(sumGridForDoseNew("f150","f151"),maxGridForDoseNew("f152","f153","f154","f155","f156"))
	_SMALL_INTESTINE_WALL = sumGridForDose(sumGridForDoseNew("f190","f191"),maxGridForDoseNew("f192","f193","f194","f195","f196"))
	_SKIN = sumGridForDose(sumGridForDoseNew("f158","f159"),maxGridForDoseNew("f160","f161","f162","f163","f164"))
	_TESTES = sumGridForDose(sumGridForDoseNew("f182","f183"),maxGridForDoseNew("f184","f185","f186","f187","f188"))
	_OVARIES = sumGridForDose(sumGridForDoseNew("f174","f175"),maxGridForDoseNew("f176","f177","f178","f179","f180"))
	_EFFECTIVE_DOSE = sumGridForDose(sumGridForDoseNew("f134","f135"),maxGridForDoseNew("f136","f137","f138","f139","f140"))
	
	wb = None
	if separate:
		wb = openpyxl.Workbook()
		ws = wb.worksheets[0]
		ws.title = u'Таблица 6'
		shift = 0
	else:
		if createNew:
			wb = openpyxl.Workbook()
			ws = wb.worksheets[0]
		else:
			wb = load_workbook(filename = oFileFolder+"/Table.xlsx")
			ws = wb.worksheets[0]

	startCol = 1; endCol = 8; sratrRow = 1+shift; endRow = 1+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A1:H1')
	ws.cell(row = sratrRow, column = startCol,  value = u'Таблица 6 – Прогнозируемые эквивалентные и эффективная дозы с уровнями доверия varSeries процентов за счет: облучения от облака и от выпадений на поверхность земли и внутреннего облучения от ингаляционного поступления радионуклидов для детей (доза за 10 суток), Зв . scenario'.replace("scenario",scenario).replace("varSeries", nameConfInt))
	
	startCol = 1; endCol = 1; sratrRow = 2+shift; endRow = 3+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A2:A3')
	ws.cell(row = sratrRow, column = startCol,  value = u'Населенные пункты (долгота; широта; расстояние, км)')
	
	startCol = 2; endCol = 8; sratrRow = 2+shift; endRow = 2+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('B2:H2')
	ws.cell(row = sratrRow, column = startCol,  value = u'Дозы на органы (дети)')
	
	writeOrgans(ws, u'красный костный мозг', u'щитовидная железа', u'легкие', u'тонкий кишечник',  u'кожа', u'гонады', u'эффективная доза')
	
	
	organs = 7
	curPos = 4
	
	for pointId, point in enumerate(pointsForAnalysis):	
		ws.cell(row = curPos+pointId*len(listOfConInt)+shift, column = 1, value = point[0]+" ("+ format(point[1], '.6f').replace('.', ',')+";"+format(point[2], '.6f').replace('.', ',')+";"+format(point[3], '.1f').replace('.', ',')+")")
		for  ConIntId, ConInt in enumerate(listOfConInt):
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=2, value=getValueForConfidence(pointId,_RED_MARROW, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=3, value=getValueForConfidence(pointId,_THYROID, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=4, value=getValueForConfidence(pointId,_LUNGS, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=5, value=getValueForConfidence(pointId,_SMALL_INTESTINE_WALL, float(ConInt)) )
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=6, value=getValueForConfidence(pointId,_SKIN, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=7, value=max(getValueForConfidence(pointId,_OVARIES, float(ConInt)),getValueForConfidence(pointId,_TESTES, float(ConInt))))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=8, value=getValueForConfidence(pointId,_EFFECTIVE_DOSE, float(ConInt)) )

	shift = shift + curPos+ConIntId+pointId*len(listOfConInt)+2
	for pointId, point in enumerate(pointsForAnalysis):
		for i in range(0, organs):
			ws.cell(row=curPos+pointId, column=2+i).number_format = '0.00E+00'
	
	if separate:		
		wb.save(oFileFolder+"/Table6.xlsx")
	else:
		wb.save(oFileFolder+"/Table.xlsx")
	
	del _RED_MARROW[:]
	del _THYROID[:]
	del _LUNGS[:]
	del _SMALL_INTESTINE_WALL[:]
	del _SKIN[:]
	del _TESTES[:]
	del _OVARIES[:]
	del _EFFECTIVE_DOSE[:]
	
	return

# Таблица 7 – Прогнозируемые дозы за счет внешнего облучения от облака и от выпадений на поверхность земли и внутреннего облучения от ингаляционного поступления радионуклидов (взрослые), Зв . scenario
def createTable7(scenario, separate, createNew):
	global nameConfInt, shift
	_THYROID_7 = sumGridForDoseNew("f200","f201","f173")
	_FOETUS_7 = sumGridForDoseNew("f3","f5","f2")
	_EFFECTIVE_DOSE_7 = sumGridForDoseNew("f198","f199","f141")
	_FOETUS_365 = sumGridForDoseNew("f3","f4","f2")
	_EFFECTIVE_DOSE_365 = sumGridForDoseNew("f204","f205","f141")
	
	wb = None
	if separate:
		wb = openpyxl.Workbook()
		ws = wb.worksheets[0]
		ws.title = u'Таблица 7'
		shift = 0
	else:
		if createNew:
			wb = openpyxl.Workbook()
			ws = wb.worksheets[0]
		else:
			wb = load_workbook(filename = oFileFolder+"/Table.xlsx")
			ws = wb.worksheets[0]
	
	startCol = 1; endCol = 6; sratrRow = 1+shift; endRow = 1+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A1:F1')
	ws.cell(row = sratrRow, column = startCol,  value = u'Таблица 7 – Прогнозируемые дозы с уровнями доверия varSeries процентов за счет внешнего облучения от облака и от выпадений на поверхность земли и внутреннего облучения от ингаляционного поступления радионуклидов (взрослые), Зв . scenario'.replace("scenario",scenario).replace("varSeries", nameConfInt))
	
	startCol = 1; sratrRow = 2+shift
	ws.cell(row = sratrRow, column = startCol,  value =u'Населенные пункты (долгота; широта; расстояние, км)')
	
	writeOrgans(ws, u'Эквивалентная доза на щитовидную железу за 7 суток', u'Эквивалентная доза на плод за 7 суток', u'Эффективная доза за 7 суток', u'Эквивалентная доза на плод за первый год', u'Эффективная доза за первый год')
	
	organs = 5
	curPos = 3
	
	for pointId, point in enumerate(pointsForAnalysis):
		ws.cell(row = curPos+pointId*len(listOfConInt)+shift, column = 1, value = point[0]+" ("+ format(point[1], '.6f').replace('.', ',')+";"+format(point[2], '.6f').replace('.', ',')+";"+format(point[3], '.1f').replace('.', ',')+")")
		for  ConIntId, ConInt in enumerate(listOfConInt):
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=2, value=getValueForConfidence(pointId,_THYROID_7, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=3, value=getValueForConfidence(pointId,_FOETUS_7, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=4, value=getValueForConfidence(pointId,_EFFECTIVE_DOSE_7, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=5, value=getValueForConfidence(pointId,_FOETUS_365, float(ConInt)) )
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=6, value=getValueForConfidence(pointId,_EFFECTIVE_DOSE_365, float(ConInt)))
	shift = shift + curPos+ConIntId+pointId*len(listOfConInt)+2
	for pointId, point in enumerate(pointsForAnalysis):
		for i in range(0, organs):
			ws.cell(row=curPos+pointId, column=2+i).number_format = '0.00E+00'

	
	if separate:		
		wb.save(oFileFolder+"/Table7.xlsx")
	else:
		wb.save(oFileFolder+"/Table.xlsx")
	
	del _THYROID_7[:]
	del _FOETUS_7[:]
	del _EFFECTIVE_DOSE_7[:]
	del _FOETUS_365[:]
	del _EFFECTIVE_DOSE_365[:]
	
	return

# Таблица 8 – Прогнозируемые дозы за счет внешнего облучения от облака и от выпадений на поверхность земли и внутреннего облучения от ингаляционного поступления радионуклидов (дети), Зв . scenario
def createTable8(scenario, separate, createNew):
	global nameConfInt, shift
	_THYROID_7 =    sumGridForDose(sumGridForDoseNew("f200","f201"),maxGridForDoseNew("f168","f169","f170","f171","f172"))
	_EFFECTIVE_DOSE_7 =  sumGridForDose(sumGridForDoseNew("f198","f199"),maxGridForDoseNew("f136","f137","f138","f139","f140"))
	_EFFECTIVE_DOSE_365 = sumGridForDose(sumGridForDoseNew("f204","f205"),maxGridForDoseNew("f136","f137","f138","f139","f140"))
	
	wb = None
	if separate:
		wb = openpyxl.Workbook()
		ws = wb.worksheets[0]
		ws.title = u'Таблица 8'
		shift = 0
	else:
		if createNew:
			wb = openpyxl.Workbook()
			ws = wb.worksheets[0]
		else:
			wb = load_workbook(filename = oFileFolder+"/Table.xlsx")
			ws = wb.worksheets[0]
	
	
	startCol = 1; endCol = 4; sratrRow = 1+shift; endRow = 1+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A1:D1')
	ws.cell(row = sratrRow, column = startCol,  value = u'Таблица 8 – Прогнозируемые дозы с уровнями доверия varSeries процентов за счет внешнего облучения от облака и от выпадений на поверхность земли и внутреннего облучения от ингаляционного поступления радионуклидов (дети), Зв . scenario'.replace("scenario",scenario).replace("varSeries", nameConfInt))
	
	startCol = 1; sratrRow = 2+shift
	ws.cell(row = sratrRow, column = startCol,  value =u'Населенные пункты (долгота; широта; расстояние, км)')
	
	writeOrgans(ws, u'Эквивалентная доза на щитовидную железу за 7 суток', u'Эффективная доза за 7 суток', u'Эффективная доза за первый год')
	
	organs = 3
	curPos = 3
	
	for pointId, point in enumerate(pointsForAnalysis):
		ws.cell(row = curPos+pointId*len(listOfConInt)+shift, column = 1, value = point[0]+" ("+ format(point[1], '.6f').replace('.', ',')+";"+format(point[2], '.6f').replace('.', ',')+";"+format(point[3], '.1f').replace('.', ',')+")")
		for  ConIntId, ConInt in enumerate(listOfConInt):
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=2, value=getValueForConfidence(pointId,_THYROID_7, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=3, value=getValueForConfidence(pointId,_EFFECTIVE_DOSE_7, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=4, value=getValueForConfidence(pointId,_EFFECTIVE_DOSE_365, float(ConInt)))
	shift = shift + curPos+ConIntId+pointId*len(listOfConInt)+2
	for pointId, point in enumerate(pointsForAnalysis):
		for i in range(0, organs):
			ws.cell(row=curPos+pointId, column=2+i).number_format = '0.00E+00'
	
	if separate:		
		wb.save(oFileFolder+"/Table8.xlsx")
	else:
		wb.save(oFileFolder+"/Table.xlsx")
	
	del _THYROID_7[:]
	del _EFFECTIVE_DOSE_7[:]
	del _EFFECTIVE_DOSE_365[:]
	
	return

# Таблица 9 – Прогнозируемые дозы за счет попадания радионуклидов внутрь организма пероральным путем, Зв . scenario 
def createTable9(createNew, scenario,milk_plants_meat = [0.0]*3):
	global nameConfInt, shift
	_FALLOUT = sumGridForDoseNew("f0")
	
	wb = None
	if separate:
		wb = openpyxl.Workbook()
		ws = wb.worksheets[0]
		ws.title = u'Таблица 9'
		shift = 0
	else:
		if createNew:
			wb = openpyxl.Workbook()
			ws = wb.worksheets[0]
		else:
			wb = load_workbook(filename = oFileFolder+"/Table.xlsx")
			ws = wb.worksheets[0]
	
	startCol = 1; endCol = 4; sratrRow = 1+shift; endRow = 1+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A1:D1')
	ws.cell(row = sratrRow, column = startCol,  value = u'Таблица 9 – Прогнозируемые дозы с уровнями доверия varSeries процентов за счет попадания радионуклидов внутрь организма пероральным путем, Зв . scenario'.replace("scenario",scenario))
	
	startCol = 1; sratrRow = 2+shift
	ws.cell(row = sratrRow, column = startCol,  value =u'Населенные пункты (долгота; широта; расстояние, км)')
	
	writeOrgans(ws, u'Эффективная доза за первый год за счет потребления молока', u'Эффективная доза за первый год за счет потребления мяса', u'Эффективная доза за первый год за счет потребления овощей')
	
	organs = 3
	curPos = 3
	
	for pointId, point in enumerate(pointsForAnalysis):
		ws['A'+str(curPos+pointId*len(listOfConInt)+shift)] = point[0]+" ("+ format(point[1], '.6f').replace('.', ',')+";"+format(point[2], '.6f').replace('.', ',')+";"+format(point[3], '.1f').replace('.', ',')+")"
		for  ConIntId, ConInt in enumerate(listOfConInt):
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=2, value=getValueForConfidence(pointId,_FALLOUT, float(ConInt))*milk_plants_meat[0])
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=3, value=getValueForConfidence(pointId,_FALLOUT, float(ConInt))*milk_plants_meat[2])
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=4, value=getValueForConfidence(pointId,_FALLOUT, float(ConInt))*milk_plants_meat[1])
	shift = shift + curPos+ConIntId+pointId*len(listOfConInt)+2
	for pointId, point in enumerate(pointsForAnalysis):
		for i in range(0, organs):
			ws.cell(row=curPos+pointId, column=2+i).number_format = '0.00E+00'
	
	if separate:		
		wb.save(oFileFolder+"/Table9.xlsx")
	else:
		wb.save(oFileFolder+"/Table.xlsx")
	
	del _FALLOUT[:]
	return

#Таблица 10 – Прогнозируемые плотность поверхностных выпадений и проинтегрированная по времени концентрация
def createTable10(scenario, separate, createNew):
	global nameConfInt, shift
	_FALLOUT = sumGridForDoseNew("f0")
	_TIC = sumGridForDoseNew("f1")
	
	wb = None
	if separate:
		wb = openpyxl.Workbook()
		ws = wb.worksheets[0]
		ws.title = u'Таблица 10'
		shift = 0
	else:
		if createNew:
			wb = openpyxl.Workbook()
			ws = wb.worksheets[0]
		else:
			wb = load_workbook(filename = oFileFolder+"/Table.xlsx")
			ws = wb.worksheets[0]
	
	startCol = 1; endCol = 3; sratrRow = 1+shift; endRow = 1+shift
	ws.merge_cells(start_row=sratrRow, start_column=startCol, end_row=endRow, end_column=endCol) #('A1:C1')
	ws.cell(row = sratrRow, column = startCol,  value = u'Таблица 10 – Прогнозируемые с уровнями доверия varSeries процентов плотность поверхностных выпадений и проинтегрированная по времени концентрация. scenario'.replace("scenario",scenario).replace("varSeries", nameConfInt))
	startCol = 1; sratrRow = 2+shift
	ws.cell(row = sratrRow, column = startCol,  value =u'Населенные пункты (долгота; широта; расстояние, км)')
	
	writeOrgans(ws, u'Плотность поверхностных выпадений, Бк/кв.м.', u'Проинтегрированная по времени концентрация, Бк·с/кб.м.')
	
	organs = 2
	curPos = 3
	for pointId, point in enumerate(pointsForAnalysis):
		ws.cell(row = curPos+pointId*len(listOfConInt)+shift, column = 1, value = point[0]+" ("+ format(point[1], '.6f').replace('.', ',')+";"+format(point[2], '.6f').replace('.', ',')+";"+format(point[3], '.1f').replace('.', ',')+")")
		for  ConIntId, ConInt in enumerate(listOfConInt):
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=2, value=getValueForConfidence(pointId,_FALLOUT, float(ConInt)))
			ws.cell(row=curPos+ConIntId+pointId*len(listOfConInt)+shift, column=3, value=getValueForConfidence(pointId,_TIC, float(ConInt)))
	shift = shift + curPos+ConIntId+pointId*len(listOfConInt)+2
	for pointId, point in enumerate(pointsForAnalysis):
		for i in range(0, organs):
			ws.cell(row=curPos+pointId, column=2+i).number_format = '0.00E+00'

	
	if separate:		
		wb.save(oFileFolder+"/Table10.xlsx")
	else:
		wb.save(oFileFolder+"/Table.xlsx")
	
	del _FALLOUT[:]
	del _TIC[:]
	return
def setDocumentStyle():
	document = Document()
	my_styles = document.styles	
	
	p_style = my_styles.add_style('name-of-p-style', WD_STYLE_TYPE.PARAGRAPH)
	p_style.base_style = my_styles['Normal']
	p_style.paragraph_format.space_before = Pt(0)
	p_style.paragraph_format.space_after = Pt(10)
	
	ch_style = my_styles.add_style('name-of-ch-style', WD_STYLE_TYPE.CHARACTER)
	ch_style.base_style = my_styles['Default Paragraph Font']
	ch_style.font.name = 'Times New Roman'
	ch_style.font.size = Pt(12)
	
	t_style = my_styles.add_style('name-of-t-style', WD_STYLE_TYPE.PARAGRAPH)
	t_style.base_style = my_styles['Normal']
	t_style.paragraph_format.space_before = Pt(0)
	t_style.paragraph_format.space_after = Pt(10)
	
	ch_style = my_styles.add_style('name-of-ch_t-style', WD_STYLE_TYPE.CHARACTER)
	ch_style.base_style = my_styles['Default Paragraph Font']
	ch_style.font.name = 'Times New Roman'
	ch_style.font.size = Pt(8)
	document.add_heading(u'Приложение 1. Расчетные значения базовых функционалов в населенных пунктах', 0)
	return document
def main(iFileFolder,  oFileFolder, oPrefixFileName, oFileType, pathToNP, listOfConInt, separate):
	mypath=os.path.dirname(os.path.realpath( __file__ ))
	os.chdir(mypath)
	global pointsForAnalysis, arrayOfResultsForPoints, nameConfInt
	pointsForAnalysis = reaNPFromFile(pathToNP)
	arrayOfResultsForPoints = [[] for x in pointsForAnalysis]
	print pointsForAnalysis
	myWorkPath = oFileFolder
	make_sure_path_exists(myWorkPath)
	prepareArrayForVarSeries()
	
	first = True
	for i in range(len(listOfConInt)):
		
		if i != len(listOfConInt)-1:
			if first:
				nameConfInt = nameConfInt + listOfConInt[i]
				first = False
			else:
				nameConfInt = nameConfInt + ", "+ listOfConInt[i]
		else:
			nameConfInt = nameConfInt + " и "+ listOfConInt[i]
	nameConfInt = nameConfInt.decode('utf-8')
	#print nameConfInt
	if oFileType == "docx":

		document = setDocumentStyle()
		
		scenario = u"Полное обесточивание без управления"
		mainTitle = u"Расчетные плотность поверхностных выпадений и проинтегрированная по времени концентрация с уровнями доверия varSeries процентов".replace("varSeries", nameConfInt)
		multiplier = 1.0
		addSingleTable(document,scenario,mainTitle,multiplier,separate, ["f0",u'Плотность поверхностных выпадений, Бк/кв.м.'],["f1",u'Проинтегрированная по времени концентрация, Бк·с/кб.м.'])		
		
		
		mainTitle = u"Расчетная эквивалентная доза облучения на плод с уровнями доверия varSeries процентов, мЗв".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addSingleTable(document,scenario,mainTitle,multiplier,separate,\
									["f2",u'Ингаляция'],\
									["f3",u'Внешнее облучение от проходящего облака'],\
									["f4",u'Внешнее облучения от поверхностных выпадений (год)'],\
									["f5",u'Внешнее облучения от поверхностных выпадений (7 суток)'],\
									["f6",u'Внешнее облучения от поверхностных выпадений (2 суток)'],\
									["f7",u'Внешнее облучения от поверхностных выпадений (10 часов)'],\
									["f8",u'Внешнее облучения от поверхностных выпадений (сутки)'],\
									["f9",u'Внешнее облучения от поверхностных выпадений (30 суток)'])		
		
		mainTitle = u"Расчетная ОБЭ - взвешенная поглощенная доза за 10 часов с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addSingleTable(document,scenario,mainTitle,multiplier,separate, \
									["f10",u'На красный костный мозг за счет внешнего облучения от проходящего облака'],\
									["f11",u'На красный костный мозг за счет внешнего облучения от поверхностных выпадений'],\
									["f12",u'На легкие за счет внешнего облучения от проходящего облака'],\
									["f13",u'На легкие за счет внешнего облучения от поверхностных выпадений'],\
									["f14",u'На яичники за счет внешнего облучения от проходящего облака'],\
									["f15",u'На яичники за счет внешнего облучения от поверхностных выпадений'],\
									["f16",u'На семенники за счет внешнего облучения от проходящего облака'],\
									["f17",u'На семенники за счет внешнего облучения от поверхностных выпадений'],\
									["f18",u'На щитовидную железу за счет внешнего облучения от проходящего облака'],\
									["f19",u'На щитовидную железу за счет внешнего облучения от поверхностных выпадений'])
		
		mainTitle = u"Расчетная ОБЭ - взвешенная поглощенная доза за 1 день на красный костный мозг за счет внешнего облучения с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addSingleTable(document,scenario,mainTitle,multiplier,separate, ["f20",u'От проходящего облака'],["f21",u'От поверхностных выпадений'])		
		
	
		mainTitle = u"Расчетная ОБЭ - взвешенная поглощенная доза за 30 суток на все тело с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f22",u'От проходящего облака'],\
									["f23",u'От поверхностных выпадений'],\
									["f24",u'На новорожденных'],\
									["f25",u'На детей 1-2 года'],\
									["f26",u'На детей 2-7 года'],\
									["f27",u'На детей 7-12 лет'],\
									["f28",u'На детей 12-17 лет'],\
									["f29",u'На взрослых'])
	
		mainTitle = u"Расчетная ОБЭ - взвешенная поглощенная доза за 30 суток на красный костный мозг с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f30",u'От проходящего облака'],\
									["f31",u'От поверхностных выпадений'],\
									["f32",u'На новорожденных'],\
									["f33",u'На детей 1-2 года'],\
									["f34",u'На детей 2-7 года'],\
									["f35",u'На детей 7-12 лет'],\
									["f36",u'На детей 12-17 лет'],\
									["f37",u'На взрослых'])
	
		mainTitle = u"Расчетная ОБЭ - взвешенная поглощенная доза за 30 суток на легкие с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f38",u'От проходящего облака'],\
									["f39",u'От поверхностных выпадений'],\
									["f40",u'На новорожденных'],\
									["f41",u'На детей 1-2 года'],\
									["f42",u'На детей 2-7 года'],\
									["f43",u'На детей 7-12 лет'],\
									["f44",u'На детей 12-17 лет'],\
									["f45",u'На взрослых'])
	
		mainTitle = u"Расчетная ОБЭ - взвешенная поглощенная доза за 30 суток на щитовидную железу с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f46",u'От проходящего облака'],\
									["f47",u'От поверхностных выпадений'],\
									["f48",u'На новорожденных'],\
									["f49",u'На детей 1-2 года'],\
									["f50",u'На детей 2-7 года'],\
									["f51",u'На детей 7-12 лет'],\
									["f52",u'На детей 12-17 лет'],\
									["f53",u'На взрослых'])
	
		mainTitle = u"Расчетная ОБЭ - взвешенная поглощенная доза за 30 суток на верхнюю часть толстого кишечника с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f54",u'От проходящего облака'],\
									["f55",u'От поверхностных выпадений'],\
									["f56",u'На новорожденных'],\
									["f57",u'На детей 1-2 года'],\
									["f58",u'На детей 2-7 года'],\
									["f59",u'На детей 7-12 лет'],\
									["f60",u'На детей 12-17 лет'],\
									["f61",u'На взрослых'])
		
		mainTitle = u"Расчетная ОБЭ - взвешенная поглощенная доза за 30 суток на нижнюю часть толстого кишечника с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f62",u'От проходящего облака'],\
									["f63",u'От поверхностных выпадений'],\
									["f64",u'На новорожденных'],\
									["f65",u'На детей 1-2 года'],\
									["f66",u'На детей 2-7 года'],\
									["f67",u'На детей 7-12 лет'],\
									["f68",u'На детей 12-17 лет'],\
									["f69",u'На взрослых'])	
	
		mainTitle = u"Расчетная поглощенная доза за 2 суток на все тело с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f70",u'От проходящего облака'],\
									["f71",u'От поверхностных выпадений'],\
									["f72",u'На новорожденных'],\
									["f73",u'На детей 1-2 года'],\
									["f74",u'На детей 2-7 года'],\
									["f75",u'На детей 7-12 лет'],\
									["f76",u'На детей 12-17 лет'],\
									["f77",u'На взрослых'])	
	
		mainTitle = u"Расчетная поглощенная доза за 2 суток на красный костный мозг с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f78",u'От проходящего облака'],\
									["f79",u'От поверхностных выпадений'],\
									["f80",u'На новорожденных'],\
									["f81",u'На детей 1-2 года'],\
									["f82",u'На детей 2-7 года'],\
									["f83",u'На детей 7-12 лет'],\
									["f84",u'На детей 12-17 лет'],\
									["f85",u'На взрослых'])	
	
		mainTitle = u"Расчетная поглощенная доза за 2 суток на легкие с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f86",u'От проходящего облака'],\
									["f87",u'От поверхностных выпадений'],\
									["f88",u'На новорожденных'],\
									["f89",u'На детей 1-2 года'],\
									["f90",u'На детей 2-7 года'],\
									["f91",u'На детей 7-12 лет'],\
									["f92",u'На детей 12-17 лет'],\
									["f93",u'На взрослых'])	
	
		mainTitle = u"Расчетная поглощенная доза за 2 суток на кожу с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f94",u'От проходящего облака'],\
									["f95",u'От поверхностных выпадений'],\
									["f96",u'На новорожденных'],\
									["f97",u'На детей 1-2 года'],\
									["f98",u'На детей 2-7 года'],\
									["f99",u'На детей 7-12 лет'],\
									["f100",u'На детей 12-17 лет'],\
									["f101",u'На взрослых'])	
	
		mainTitle = u"Расчетная поглощенная доза за 2 суток на щитовидную железу с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f102",u'От проходящего облака'],\
									["f103",u'От поверхностных выпадений'],\
									["f104",u'На новорожденных'],\
									["f105",u'На детей 1-2 года'],\
									["f106",u'На детей 2-7 года'],\
									["f107",u'На детей 7-12 лет'],\
									["f108",u'На детей 12-17 лет'],\
									["f109",u'На взрослых'])
	
		mainTitle = u"Расчетная поглощенная доза за 2 суток на яичники с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f110",u'От проходящего облака'],\
									["f111",u'От поверхностных выпадений'],\
									["f112",u'На новорожденных'],\
									["f113",u'На детей 1-2 года'],\
									["f114",u'На детей 2-7 года'],\
									["f115",u'На детей 7-12 лет'],\
									["f116",u'На детей 12-17 лет'],\
									["f117",u'На взрослых'])
	
		mainTitle = u"Расчетная поглощенная доза за 2 суток на семенники с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f118",u'От проходящего облака'],\
									["f119",u'От поверхностных выпадений'],\
									["f120",u'На новорожденных'],\
									["f121",u'На детей 1-2 года'],\
									["f122",u'На детей 2-7 года'],\
									["f123",u'На детей 7-12 лет'],\
									["f124",u'На детей 12-17 лет'],\
									["f125",u'На взрослых'])
	
		mainTitle = u"Расчетная поглощенная доза за 2 суток на тонкий кишечник с уровнями доверия varSeries процентов, мГр".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f126",u'От проходящего облака'],\
									["f127",u'От поверхностных выпадений'],\
									["f128",u'На новорожденных'],\
									["f129",u'На детей 1-2 года'],\
									["f130",u'На детей 2-7 года'],\
									["f131",u'На детей 7-12 лет'],\
									["f132",u'На детей 12-17 лет'],\
									["f133",u'На взрослых'])
	
		mainTitle = u"Расчетная эффективная доза за 10 суток с уровнями доверия varSeries процентов, мЗв".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f134",u'От проходящего облака'],\
									["f135",u'От поверхностных выпадений'],\
									["f136",u'На новорожденных'],\
									["f137",u'На детей 1-2 года'],\
									["f138",u'На детей 2-7 года'],\
									["f139",u'На детей 7-12 лет'],\
									["f140",u'На детей 12-17 лет'],\
									["f141",u'На взрослых'])
	
		mainTitle = u"Расчетная эквивалентная доза за 10 суток на красный костный мозг с уровнями доверия varSeries процентов, мЗв".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f142",u'От проходящего облака'],\
									["f143",u'От поверхностных выпадений'],\
									["f144",u'На новорожденных'],\
									["f145",u'На детей 1-2 года'],\
									["f146",u'На детей 2-7 года'],\
									["f147",u'На детей 7-12 лет'],\
									["f148",u'На детей 12-17 лет'],\
									["f149",u'На взрослых'])
	
		mainTitle = u"Расчетная эквивалентная доза за 10 суток на легкие с уровнями доверия varSeries процентов, мЗв".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f150",u'От проходящего облака'],\
									["f151",u'От поверхностных выпадений'],\
									["f152",u'На новорожденных'],\
									["f153",u'На детей 1-2 года'],\
									["f154",u'На детей 2-7 года'],\
									["f155",u'На детей 7-12 лет'],\
									["f156",u'На детей 12-17 лет'],\
									["f157",u'На взрослых'])
		
		mainTitle = u"Расчетная эквивалентная доза за 10 суток на кожу с уровнями доверия varSeries процентов, мЗв".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f158",u'От проходящего облака'],\
									["f159",u'От поверхностных выпадений'],\
									["f160",u'На новорожденных'],\
									["f161",u'На детей 1-2 года'],\
									["f162",u'На детей 2-7 года'],\
									["f163",u'На детей 7-12 лет'],\
									["f164",u'На детей 12-17 лет'],\
									["f165",u'На взрослых'])
	
		mainTitle = u"Расчетная эквивалентная доза за 10 суток на щитовидную железу с уровнями доверия varSeries процентов, мЗв".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f166",u'От проходящего облака'],\
									["f167",u'От поверхностных выпадений'],\
									["f168",u'На новорожденных'],\
									["f169",u'На детей 1-2 года'],\
									["f170",u'На детей 2-7 года'],\
									["f171",u'На детей 7-12 лет'],\
									["f172",u'На детей 12-17 лет'],\
									["f173",u'На взрослых'])
	
		mainTitle = u"Расчетная эквивалентная доза за 10 суток на яичники с уровнями доверия varSeries процентов, мЗв".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f174",u'От проходящего облака'],\
									["f175",u'От поверхностных выпадений'],\
									["f176",u'На новорожденных'],\
									["f177",u'На детей 1-2 года'],\
									["f178",u'На детей 2-7 года'],\
									["f179",u'На детей 7-12 лет'],\
									["f180",u'На детей 12-17 лет'],\
									["f181",u'На взрослых'])
	
		mainTitle = u"Расчетная эквивалентная доза за 10 суток на семенники с уровнями доверия varSeries процентов, мЗв".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f182",u'От проходящего облака'],\
									["f183",u'От поверхностных выпадений'],\
									["f184",u'На новорожденных'],\
									["f185",u'На детей 1-2 года'],\
									["f186",u'На детей 2-7 года'],\
									["f187",u'На детей 7-12 лет'],\
									["f188",u'На детей 12-17 лет'],\
									["f189",u'На взрослых'])
	
		mainTitle = u"Расчетная эквивалентная доза за 10 суток на тонкий кишечник с уровнями доверия varSeries процентов, мЗв".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addComplexTable(document,scenario,mainTitle,multiplier,separate, \
									["f190",u'От проходящего облака'],\
									["f191",u'От поверхностных выпадений'],\
									["f192",u'На новорожденных'],\
									["f193",u'На детей 1-2 года'],\
									["f194",u'На детей 2-7 года'],\
									["f195",u'На детей 7-12 лет'],\
									["f196",u'На детей 12-17 лет'],\
									["f197",u'На взрослых'])
	
		mainTitle = u"Расчетные эффективная и эквивалентная дозы за счет внешнего облучения с уровнями доверия varSeries процентов, мЗв".replace("varSeries", nameConfInt)
		multiplier = 1000.0
		addSingleTable(document,scenario,mainTitle,multiplier,separate, \
									["f198",u'За 7 суток за счет внешнего облучения от проходящего облака'],\
									["f199",u'За 7 суток за счет внешнего облучения от поверхностных выпадений'],\
									["f200",u'За 7 суток на щитовидную железу за счет внешнего облучения от проходящего облака'],\
									["f201",u'За 7 суток на щитовидную железу за счет внешнего облучения от поверхностных выпадений'],\
									["f202",u'За 30 суток за счет внешнего облучения от проходящего облака'],\
									["f203",u'За 30 суток за счет внешнего облучения от поверхностных выпадений'],\
									["f204",u'За год за счет внешнего облучения от проходящего облака'],\
									["f205",u'За год за счет внешнего облучения от поверхностных выпадений'])

		if separate == False:
			document.save(oFileFolder+'/appendix1.docx')
	elif oFileType == "xlsx":
		scenario = u"Полное обесточивание без управления"
		createNew = True
		createTable1(scenario, separate, createNew)
		createNew = False
		createTable2(scenario, separate, createNew)
		createTable3(scenario, separate, createNew)
		createTable4(scenario, separate, createNew)
		createTable5(scenario, separate, createNew)
		createTable6(scenario, separate, createNew)
		createTable7(scenario, separate, createNew)
		createTable8(scenario, separate, createNew)
		createTable10(scenario, separate, createNew)
		createTable9(createNew, scenario,milk_plants_meat = [2.82E-08, 1.04E-08, 1.57E-08])
	print("Ok!")
	return

if __name__ == "__main__":
	mypath=os.path.dirname(os.path.realpath( __file__ ))
	os.chdir(mypath)
	if len(sys.argv) == 1: sys.argv[1:] = ["-h"]
	parser = argparse.ArgumentParser()
	parser._action_groups.pop()
	required = parser.add_argument_group('required arguments')
	optional = parser.add_argument_group('optional arguments')
	
	required.add_argument('-ifwr','--inputFolderWithResults', action = 'store', dest='iFolderWithResults',  type=str, help='Path to folder with bin', required=True)
	required.add_argument('-ptnp','--pathToNP', action=arghelper.checkAndPrepareF({'dat'}), dest='pathToNP', metavar='FILE', type=lambda x: arghelper.is_valid_file(parser, x), help='path to the file np.dat', required=True)
	required.add_argument('-oft','--oftype', choices=['docx', 'xlsx'], dest='oFileType',  type=str, help='type of the output file with nostra results', required=True)
	optional.add_argument ('-separate', action='store_const', dest='separate',help='Use this flag for separate documents (Default True)',  const=True)
	
	optional.add_argument('-off','--offolder', action = 'store', dest='oFileFolder',metavar='OUTPUT_FOLDER', type=lambda x: arghelper.is_valid_directory(parser, x), help='path to the output folder for output files', required=False, default=mypath)
	optional.add_argument('-ofnf','--ofnewfolder', action = 'store', dest='oFileNewFolder',metavar='NEW_FOLDER', type=str, help='path to the output new folder for output files', required=False, default="")
	optional.add_argument('-opfn','--opfName', action = 'store', dest='oPrefixFileName',  type=str, help='Prefix to name to output file', required=False, default="")
	optional.add_argument ('--listConfInt', nargs='+', dest='listOfConInt',help='Write space separated values of confidence intervals', default=['50', '95', '99.5'])
	args_new = parser.parse_args()
	
	oFileFolder = args_new.oFileFolder
	if args_new.oFileNewFolder != '':
		if os.path.exists(args_new.oFileNewFolder) == False:
			make_sure_path_exists(args_new.oFileNewFolder)
		oFileFolder = args_new.oFileNewFolder.strip()
	pathToNP = args_new.pathToNP
	iFileFolder = args_new.iFolderWithResults.strip()
	oPrefixFileName = args_new.oPrefixFileName.strip()
	oFileType = args_new.oFileType.strip()
	listOfConInt = args_new.listOfConInt
	separate = args_new.separate
	if oPrefixFileName != "":
		oPrefixFileName = "{}".format((str(oPrefixFileName)+"_"))
	oFilePath = oFileFolder
	err = False
	for i in listOfConInt:
		try:
			float(i)
		except ValueError:
			err = True
	if separate == None:
		separate = False
	if err == False:
		sys.exit(main(iFileFolder,  oFileFolder, oPrefixFileName, oFileType,pathToNP,  listOfConInt, separate))
	else:
		print "Error with confidence intervals"


